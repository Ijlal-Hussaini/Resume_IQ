import io
import re
import logging
from typing import Tuple, Optional

logger = logging.getLogger("resumeiq.parser")


class DocumentParser:
    """Extracts and normalizes text from PDF, DOCX, TXT, and Image files."""

    @staticmethod
    def parse_file(filename: str, content: bytes) -> Tuple[str, str]:
        """
        Parses binary content based on file extension.
        Returns (extracted_text, file_type)
        """
        lower_name = filename.lower()
        
        if lower_name.endswith(".pdf"):
            text = DocumentParser._parse_pdf(content)
            return DocumentParser._normalize_text(text), "pdf"
        elif lower_name.endswith(".docx") or lower_name.endswith(".doc"):
            text = DocumentParser._parse_docx(content)
            return DocumentParser._normalize_text(text), "docx"
        elif lower_name.endswith(".txt") or lower_name.endswith(".md"):
            text = DocumentParser._parse_txt(content)
            return DocumentParser._normalize_text(text), "txt"
        elif any(lower_name.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".webp", ".tiff"]):
            text = DocumentParser._parse_image_ocr(content)
            return DocumentParser._normalize_text(text), "image_ocr"
        else:
            # Attempt generic text decode
            text = DocumentParser._parse_txt(content)
            return DocumentParser._normalize_text(text), "unknown"

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """Extracts text using pdfplumber with pypdf fallback."""
        text_parts = []

        # Try pdfplumber first for high layout accuracy
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(layout=True) or page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            if text_parts:
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}. Falling back to pypdf...")

        # Fallback: pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            if text_parts:
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"pypdf extraction failed: {e}")

        # If both fail or PDF is a scanned image, try OCR
        return DocumentParser._parse_image_ocr(content)

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """Extracts text from DOCX files."""
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text inside tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_cells:
                        paragraphs.append(" | ".join(row_cells))
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"docx extraction failed: {e}")
            return DocumentParser._parse_txt(content)

    @staticmethod
    def _parse_txt(content: bytes) -> str:
        """Decodes raw text with fallback encodings."""
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_image_ocr(content: bytes) -> str:
        """Attempts OCR on image or scanned PDF bytes."""
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(io.BytesIO(content))
            return pytesseract.image_to_string(img)
        except Exception as e:
            logger.info(f"pytesseract OCR not available or failed: {e}")

        try:
            import easyocr
            reader = easyocr.Reader(['en'], gpu=False)
            results = reader.readtext(content)
            return "\n".join([res[1] for res in results])
        except Exception as e:
            logger.info(f"easyocr not available: {e}")

        return "Could not perform OCR on image. Text extraction required."

    @staticmethod
    def _normalize_text(text: str) -> str:
        """Cleans ligatures, unusual whitespace, and normalizes headers."""
        if not text:
            return ""

        # Normalize common unicode ligatures
        ligatures = {
            "ﬁ": "fi", "ﬂ": "fl", "ﬀ": "ff", "ﬃ": "ffi", "ﬄ": "ffl",
            "\u2013": "-", "\u2014": "-", "\u2018": "'", "\u2019": "'",
            "\u201c": '"', "\u201d": '"', "\u2022": "*", "\u25cf": "*",
            "\u25aa": "*", "\u25cb": "*", "\u00a0": " "
        }
        for k, v in ligatures.items():
            text = text.replace(k, v)

        # Normalize carriage returns and multiple blanks
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Collapse >3 consecutive newlines into 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Collapse multiple spaces
        text = re.sub(r"[ \t]{2,}", " ", text)

        return text.strip()


document_parser = DocumentParser()
